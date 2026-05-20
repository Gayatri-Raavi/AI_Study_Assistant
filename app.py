from __future__ import annotations

import hashlib
import io
import json
import os
import re
import secrets
import shutil
import textwrap
import uuid
from collections import Counter
from contextlib import contextmanager
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import streamlit as st
import streamlit.components.v1 as components
from pypdf import PdfReader
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from sqlalchemy import ForeignKey, Integer, String, Text, create_engine, inspect, select, text
from sqlalchemy.orm import DeclarativeBase, Mapped, Session, mapped_column, relationship, sessionmaker

try:
    from docx import Document as DocxFile
except ModuleNotFoundError:
    DocxFile = None

try:
    from PIL import Image, ImageEnhance, ImageFilter, ImageOps
except ModuleNotFoundError:
    Image = None
    ImageEnhance = None
    ImageFilter = None
    ImageOps = None

try:
    import pytesseract
except ModuleNotFoundError:
    pytesseract = None

try:
    import speech_recognition as sr
except ModuleNotFoundError:
    sr = None


APP_TITLE = "SmartPDF AI Assistant"
BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
VECTOR_DIR = DATA_DIR / "vectors"
CLOUD_DIR = Path(os.getenv("CLOUD_STORAGE_DIR", (DATA_DIR / "cloud").as_posix()))
DATABASE_URL = os.getenv("DATABASE_URL", f"sqlite:///{(DATA_DIR / 'smartpdf.db').as_posix()}")
LOCAL_PROVIDER = "local"
DEFAULT_PROVIDER = LOCAL_PROVIDER
DEFAULT_EMBEDDING_MODEL = "offline-keyword-search"
CHUNK_SIZE = 1200
CHUNK_OVERLAP = 180
MAX_CONTEXT_CHUNKS = 6
MIN_MATCH_SCORE = 1
INFO_NOT_FOUND = "Information not found in uploaded document"
MAX_STUDY_SENTENCES = 8
IMAGE_UPLOAD_TYPES = ["png", "jpg", "jpeg", "webp"]
STOP_WORDS = {
    "a", "an", "and", "are", "as", "at", "be", "by", "do", "for", "from", "how",
    "i", "in", "is", "it", "of", "on", "or", "that", "the", "this", "to", "was",
    "what", "when", "where", "which", "who", "why", "with", "your",
}
LANGUAGE_CONFIG = {
    "English": {"code": "en", "speech": "en-US"},
    "Telugu": {"code": "te", "speech": "te-IN"},
    "Hindi": {"code": "hi", "speech": "hi-IN"},
}
TEXTS = {
    "language": {"en": "Language", "te": "భాష", "hi": "भाषा"},
    "answer_mode": {"en": "Answer mode: Offline PDF search", "te": "సమాధాన విధానం: ఆఫ్‌లైన్ PDF శోధన", "hi": "उत्तर मोड: ऑफलाइन PDF खोज"},
    "local_run": {"en": "Runs locally with optional cloud mirror", "te": "ఐచ్ఛిక క్లౌడ్ మిర్రర్‌తో లోకల్‌గా నడుస్తుంది", "hi": "वैकल्पिक क्लाउड मिरर के साथ लोकल रूप से चलता है"},
    "study_tools": {"en": "Study Tools", "te": "స్టడీ టూల్స్", "hi": "स्टडी टूल्स"},
    "smart_search": {"en": "Smart Search", "te": "స్మార్ట్ సెర్చ్", "hi": "स्मार्ट सर्च"},
    "dashboard": {"en": "Dashboard", "te": "డ్యాష్‌బోర్డ్", "hi": "डैशबोर्ड"},
    "chat": {"en": "Chat", "te": "చాట్", "hi": "चैट"},
    "summary": {"en": "PDF Summary", "te": "PDF సారాంశం", "hi": "PDF सारांश"},
    "questions": {"en": "Important Questions", "te": "ముఖ్యమైన ప్రశ్నలు", "hi": "महत्वपूर्ण प्रश्न"},
    "mcq": {"en": "MCQ Generator", "te": "MCQ జనరేటర్", "hi": "MCQ जनरेटर"},
    "notes": {"en": "Smart Notes", "te": "స్మార్ట్ నోట్స్", "hi": "स्मार्ट नोट्स"},
    "export": {"en": "Export Chat", "te": "చాట్ ఎగుమతి", "hi": "चैट एक्सपोर्ट"},
    "page_search_title": {"en": "Page-Based Smart Search", "te": "పేజీ ఆధారిత స్మార్ట్ సెర్చ్", "hi": "पेज-आधारित स्मार्ट खोज"},
    "page_search_help": {"en": "Search the selected PDFs and see matching page numbers.", "te": "ఎంచుకున్న PDFలలో శోధించి సరిపోయే పేజీ నంబర్లను చూడండి.", "hi": "चयनित PDF में खोजें और मिलते हुए पेज नंबर देखें।"},
    "search_button": {"en": "Search Pages", "te": "పేజీలు శోధించండి", "hi": "पेज खोजें"},
    "search_query": {"en": "Search inside selected PDFs", "te": "ఎంచుకున్న PDFలలో శోధించండి", "hi": "चयनित PDF में खोजें"},
    "not_found": {"en": INFO_NOT_FOUND, "te": "అప్‌లోడ్ చేసిన డాక్యుమెంట్‌లో సమాచారం దొరకలేదు", "hi": "अपलोड किए गए दस्तावेज़ में जानकारी नहीं मिली"},
    "voice_ready": {"en": "Voice draft ready", "te": "వాయిస్ డ్రాఫ్ట్ సిద్ధంగా ఉంది", "hi": "वॉइस ड्राफ्ट तैयार है"},
    "voice_output": {"en": "Voice output", "te": "Voice output", "hi": "Voice output"},
    "voice_output_help": {"en": "Play the latest AI answer with instant pause, resume, and stop controls.", "te": "Play the latest AI answer with instant pause, resume, and stop controls.", "hi": "Play the latest AI answer with instant pause, resume, and stop controls."},
    "voice_output_empty": {"en": "Ask a question first to generate an answer for voice playback.", "te": "Ask a question first to generate an answer for voice playback.", "hi": "Ask a question first to generate an answer for voice playback."},
    "cloud": {"en": "Cloud PDF Storage", "te": "క్లౌడ్ PDF నిల్వ", "hi": "क्लाउड PDF स्टोरेज"},
    "analytics": {"en": "Analytics", "te": "అనలిటిక్స్", "hi": "एनालिटिक्स"},
    "clear_current_chat": {"en": "Clear current chat", "te": "Clear current chat", "hi": "Clear current chat"},
    "clear_all_chats": {"en": "Clear all chats", "te": "Clear all chats", "hi": "Clear all chats"},
}


class Base(DeclarativeBase):
    pass


class User(Base):
    __tablename__ = "users"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    email: Mapped[str] = mapped_column(String(255), unique=True, nullable=False)
    password_hash: Mapped[str] = mapped_column(String(255), nullable=False)
    salt: Mapped[str] = mapped_column(String(64), nullable=False)
    created_at: Mapped[str] = mapped_column(String(64), nullable=False)
    theme_preference: Mapped[str] = mapped_column(String(20), default="dark", nullable=False)
    preferred_provider: Mapped[str] = mapped_column(String(20), default=DEFAULT_PROVIDER, nullable=False)

    documents: Mapped[list["Document"]] = relationship(back_populates="user", cascade="all, delete-orphan")
    chats: Mapped[list["Chat"]] = relationship(back_populates="user", cascade="all, delete-orphan")


class Document(Base):
    __tablename__ = "documents"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), nullable=False)
    original_name: Mapped[str] = mapped_column(String(255), nullable=False)
    saved_name: Mapped[str] = mapped_column(String(255), nullable=False)
    file_path: Mapped[str] = mapped_column(Text, nullable=False)
    extracted_text: Mapped[str] = mapped_column(Text, nullable=False)
    chunks_path: Mapped[str] = mapped_column(Text, nullable=False)
    embeddings_path: Mapped[str] = mapped_column(Text, nullable=False)
    index_path: Mapped[str] = mapped_column(Text, nullable=False)
    chunk_count: Mapped[int] = mapped_column(Integer, nullable=False)
    created_at: Mapped[str] = mapped_column(String(64), nullable=False)
    file_hash: Mapped[str] = mapped_column(String(64), default="", nullable=False)
    page_count: Mapped[int] = mapped_column(Integer, default=0, nullable=False)
    embedding_provider: Mapped[str] = mapped_column(String(20), default=LOCAL_PROVIDER, nullable=False)
    embedding_model: Mapped[str] = mapped_column(String(120), default="", nullable=False)

    user: Mapped[User] = relationship(back_populates="documents")


class Chat(Base):
    __tablename__ = "chats"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    user_id: Mapped[int] = mapped_column(ForeignKey("users.id"), nullable=False)
    title: Mapped[str] = mapped_column(String(255), nullable=False)
    selected_document_ids: Mapped[str] = mapped_column(Text, nullable=False)
    created_at: Mapped[str] = mapped_column(String(64), nullable=False)
    updated_at: Mapped[str] = mapped_column(String(64), nullable=False)
    provider: Mapped[str] = mapped_column(String(20), default=LOCAL_PROVIDER, nullable=False)

    user: Mapped[User] = relationship(back_populates="chats")
    messages: Mapped[list["Message"]] = relationship(back_populates="chat", cascade="all, delete-orphan")


class Message(Base):
    __tablename__ = "messages"

    id: Mapped[int] = mapped_column(Integer, primary_key=True, autoincrement=True)
    chat_id: Mapped[int] = mapped_column(ForeignKey("chats.id"), nullable=False)
    role: Mapped[str] = mapped_column(String(20), nullable=False)
    content: Mapped[str] = mapped_column(Text, nullable=False)
    created_at: Mapped[str] = mapped_column(String(64), nullable=False)
    metadata_json: Mapped[str] = mapped_column(Text, default="{}", nullable=False)

    chat: Mapped[Chat] = relationship(back_populates="messages")


def ensure_storage() -> None:
    for folder in (DATA_DIR, UPLOAD_DIR, VECTOR_DIR, CLOUD_DIR):
        folder.mkdir(parents=True, exist_ok=True)


ensure_storage()
connect_args = {"check_same_thread": False} if DATABASE_URL.startswith("sqlite") else {}
ENGINE = create_engine(DATABASE_URL, future=True, connect_args=connect_args)
SessionLocal = sessionmaker(bind=ENGINE, autoflush=False, autocommit=False, future=True)


def utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


@contextmanager
def db_session() -> Iterable[Session]:
    session = SessionLocal()
    try:
        yield session
        session.commit()
    except Exception:
        session.rollback()
        raise
    finally:
        session.close()


def init_db() -> None:
    Base.metadata.create_all(ENGINE)
    migrate_legacy_schema()


def add_column_if_missing(table_name: str, column_name: str, column_sql: str) -> None:
    inspector = inspect(ENGINE)
    columns = {column["name"] for column in inspector.get_columns(table_name)}
    if column_name in columns:
        return
    with ENGINE.begin() as connection:
        connection.execute(text(f"ALTER TABLE {table_name} ADD COLUMN {column_name} {column_sql}"))


def migrate_legacy_schema() -> None:
    try:
        add_column_if_missing("users", "theme_preference", "TEXT NOT NULL DEFAULT 'dark'")
        add_column_if_missing("users", "preferred_provider", f"TEXT NOT NULL DEFAULT '{DEFAULT_PROVIDER}'")
        add_column_if_missing("documents", "file_hash", "TEXT NOT NULL DEFAULT ''")
        add_column_if_missing("documents", "page_count", "INTEGER NOT NULL DEFAULT 0")
        add_column_if_missing("documents", "embedding_provider", f"TEXT NOT NULL DEFAULT '{LOCAL_PROVIDER}'")
        add_column_if_missing("documents", "embedding_model", "TEXT NOT NULL DEFAULT ''")
        add_column_if_missing("chats", "provider", f"TEXT NOT NULL DEFAULT '{LOCAL_PROVIDER}'")
        add_column_if_missing("messages", "metadata_json", "TEXT NOT NULL DEFAULT '{}'")
    except Exception:
        # Fresh databases or provider-specific limitations should not block the app.
        pass


def hash_password(password: str, salt: str) -> str:
    return hashlib.pbkdf2_hmac(
        "sha256",
        password.encode("utf-8"),
        salt.encode("utf-8"),
        120000,
    ).hex()


def create_user(email: str, password: str) -> tuple[bool, str]:
    email = email.strip().lower()
    if not re.fullmatch(r"[^@\s]+@[^@\s]+\.[^@\s]+", email):
        return False, "Enter a valid email address."
    if len(password) < 8:
        return False, "Password must be at least 8 characters."

    salt = secrets.token_hex(16)
    password_hash = hash_password(password, salt)

    with db_session() as session:
        existing = session.scalar(select(User).where(User.email == email))
        if existing:
            return False, "An account with that email already exists."
        session.add(
            User(
                email=email,
                password_hash=password_hash,
                salt=salt,
                created_at=utc_now(),
                theme_preference="dark",
                preferred_provider=DEFAULT_PROVIDER,
            )
        )
    return True, "Account created successfully. Please sign in."


def authenticate_user(email: str, password: str) -> dict[str, Any] | None:
    with db_session() as session:
        user = session.scalar(select(User).where(User.email == email.strip().lower()))
        if not user:
            return None
        attempted_hash = hash_password(password, user.salt)
        if not secrets.compare_digest(attempted_hash, user.password_hash):
            return None
        return {
            "id": user.id,
            "email": user.email,
            "theme_preference": user.theme_preference,
            "preferred_provider": user.preferred_provider,
        }


def update_user_preferences(user_id: int, theme: str | None = None, provider: str | None = None) -> dict[str, Any]:
    with db_session() as session:
        user = session.get(User, user_id)
        if not user:
            raise ValueError("User not found.")
        if theme:
            user.theme_preference = theme
        if provider:
            user.preferred_provider = provider
        session.add(user)
        session.flush()
        return {
            "id": user.id,
            "email": user.email,
            "theme_preference": user.theme_preference,
            "preferred_provider": user.preferred_provider,
        }


def get_user_documents(user_id: int) -> list[dict[str, Any]]:
    with db_session() as session:
        documents = session.scalars(select(Document).where(Document.user_id == user_id).order_by(Document.created_at.desc())).all()
        return [serialize_document(document) for document in documents]


def get_user_chats(user_id: int) -> list[dict[str, Any]]:
    with db_session() as session:
        chats = session.scalars(select(Chat).where(Chat.user_id == user_id).order_by(Chat.updated_at.desc())).all()
        return [serialize_chat(chat) for chat in chats]


def get_chat(chat_id: int, user_id: int) -> dict[str, Any] | None:
    with db_session() as session:
        chat = session.scalar(select(Chat).where(Chat.id == chat_id, Chat.user_id == user_id))
        return serialize_chat(chat) if chat else None


def get_messages(chat_id: int) -> list[dict[str, Any]]:
    with db_session() as session:
        messages = session.scalars(select(Message).where(Message.chat_id == chat_id).order_by(Message.created_at.asc(), Message.id.asc())).all()
        return [serialize_message(message) for message in messages]


def create_chat(user_id: int, title: str, document_ids: list[int], provider: str) -> int:
    now = utc_now()
    with db_session() as session:
        chat = Chat(
            user_id=user_id,
            title=title,
            selected_document_ids=json.dumps(document_ids),
            created_at=now,
            updated_at=now,
            provider=provider,
        )
        session.add(chat)
        session.flush()
        return int(chat.id)


def update_chat_documents(chat_id: int, document_ids: list[int]) -> None:
    with db_session() as session:
        chat = session.get(Chat, chat_id)
        if not chat:
            return
        chat.selected_document_ids = json.dumps(document_ids)
        chat.updated_at = utc_now()
        session.add(chat)


def add_message(chat_id: int, role: str, content: str, metadata: dict[str, Any] | None = None) -> None:
    with db_session() as session:
        message = Message(
            chat_id=chat_id,
            role=role,
            content=content,
            created_at=utc_now(),
            metadata_json=json.dumps(metadata or {}, ensure_ascii=True),
        )
        chat = session.get(Chat, chat_id)
        session.add(message)
        if chat:
            chat.updated_at = utc_now()
            session.add(chat)


def delete_chat(chat_id: int, user_id: int) -> bool:
    with db_session() as session:
        chat = session.scalar(select(Chat).where(Chat.id == chat_id, Chat.user_id == user_id))
        if not chat:
            return False
        session.delete(chat)
        return True


def delete_all_chats(user_id: int) -> int:
    with db_session() as session:
        chats = session.scalars(select(Chat).where(Chat.user_id == user_id)).all()
        for chat in chats:
            session.delete(chat)
        return len(chats)


def serialize_document(document: Document) -> dict[str, Any]:
    return {
        "id": document.id,
        "user_id": document.user_id,
        "original_name": document.original_name,
        "saved_name": document.saved_name,
        "file_path": document.file_path,
        "extracted_text": document.extracted_text,
        "chunks_path": document.chunks_path,
        "chunk_count": document.chunk_count,
        "created_at": document.created_at,
        "file_hash": document.file_hash,
        "page_count": document.page_count,
    }


def serialize_chat(chat: Chat) -> dict[str, Any]:
    return {
        "id": chat.id,
        "user_id": chat.user_id,
        "title": chat.title,
        "selected_document_ids": chat.selected_document_ids,
        "created_at": chat.created_at,
        "updated_at": chat.updated_at,
        "provider": chat.provider,
    }


def serialize_message(message: Message) -> dict[str, Any]:
    try:
        metadata = json.loads(message.metadata_json or "{}")
    except json.JSONDecodeError:
        metadata = {}
    return {
        "id": message.id,
        "chat_id": message.chat_id,
        "role": message.role,
        "content": message.content,
        "created_at": message.created_at,
        "metadata": metadata,
    }


def build_chat_title(question: str, selected_docs: list[dict[str, Any]]) -> str:
    if question.strip():
        return textwrap.shorten(question.strip(), width=48, placeholder="...")
    if selected_docs:
        names = ", ".join(doc["original_name"] for doc in selected_docs[:2])
        return textwrap.shorten(names, width=48, placeholder="...")
    return "New chat"


def clean_filename(name: str) -> str:
    sanitized = re.sub(r"[^A-Za-z0-9._-]+", "_", name)
    return sanitized.strip("._") or "document.pdf"


class InMemoryUpload:
    def __init__(self, name: str, file_bytes: bytes) -> None:
        self.name = name
        self._file_bytes = file_bytes

    def getvalue(self) -> bytes:
        return self._file_bytes


def compute_file_hash(file_bytes: bytes) -> str:
    return hashlib.sha256(file_bytes).hexdigest()


def configure_ocr_engine() -> str | None:
    configured = os.getenv("TESSERACT_CMD", "").strip()
    tesseract_path = configured or shutil.which("tesseract")
    if pytesseract is not None and tesseract_path:
        pytesseract.pytesseract.tesseract_cmd = tesseract_path
    return tesseract_path


def ocr_is_available() -> bool:
    return Image is not None and pytesseract is not None and bool(configure_ocr_engine())


def enhance_scan_image(image: Any) -> Any:
    if ImageOps is None or ImageEnhance is None or ImageFilter is None:
        raise RuntimeError("Image enhancement requires Pillow.")
    normalized = ImageOps.exif_transpose(image)
    grayscale = ImageOps.grayscale(normalized)
    grayscale = ImageOps.autocontrast(grayscale)
    grayscale = ImageEnhance.Contrast(grayscale).enhance(1.35)
    grayscale = grayscale.filter(ImageFilter.SHARPEN)
    return grayscale.convert("RGB")


def read_image_bytes(image_bytes: bytes) -> Any:
    if Image is None:
        raise RuntimeError("Image processing requires Pillow.")
    with Image.open(io.BytesIO(image_bytes)) as raw_image:
        prepared = enhance_scan_image(raw_image)
    return prepared


def ocr_image_to_text(image: Any) -> str:
    if not ocr_is_available():
        raise RuntimeError("OCR requires `pytesseract` and the Tesseract OCR executable installed on this machine.")
    text_value = pytesseract.image_to_string(image, config="--psm 6")
    return re.sub(r"\s+\n", "\n", text_value).strip()


def pdf_page_to_ocr_text(page: Any) -> str:
    if not ocr_is_available():
        return ""
    page_texts: list[str] = []
    for image_file in getattr(page, "images", []) or []:
        image_bytes = getattr(image_file, "data", b"")
        if not image_bytes:
            continue
        try:
            page_texts.append(ocr_image_to_text(read_image_bytes(image_bytes)))
        except Exception:
            continue
    return "\n".join(text for text in page_texts if text.strip()).strip()


def extract_text_from_pdf(file_bytes: bytes) -> tuple[str, int]:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(page_text.strip())
    return "\n\n".join(pages), len(reader.pages)


def extract_page_records(file_bytes: bytes) -> tuple[str, int, list[dict[str, Any]]]:
    reader = PdfReader(io.BytesIO(file_bytes))
    page_records: list[dict[str, Any]] = []
    pages: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if not page_text:
            page_text = pdf_page_to_ocr_text(page)
        if page_text:
            pages.append(page_text)
            page_records.append({"page": page_number, "text": page_text})
    return "\n\n".join(pages), len(reader.pages), page_records


def images_to_pdf_bytes(images: list[Any]) -> bytes:
    if not images:
        raise RuntimeError("Choose at least one scanned page.")
    rgb_images = [image.convert("RGB") for image in images]
    buffer = io.BytesIO()
    rgb_images[0].save(buffer, format="PDF", save_all=True, append_images=rgb_images[1:])
    return buffer.getvalue()


def build_scanned_document(
    camera_pages: list[dict[str, Any]],
    gallery_files: list[Any],
    requested_name: str,
) -> tuple[InMemoryUpload, str, int, list[dict[str, Any]], list[bytes]]:
    pages: list[dict[str, Any]] = list(camera_pages)
    for uploaded_image in gallery_files:
        pages.append({"name": uploaded_image.name, "bytes": uploaded_image.getvalue()})

    if not pages:
        raise RuntimeError("Add at least one camera scan or gallery image before processing.")

    enhanced_images: list[Any] = []
    page_records: list[dict[str, Any]] = []
    preview_images: list[bytes] = []
    extracted_pages: list[str] = []

    for page_number, page in enumerate(pages, start=1):
        enhanced_image = read_image_bytes(page["bytes"])
        enhanced_images.append(enhanced_image)
        preview_buffer = io.BytesIO()
        enhanced_image.save(preview_buffer, format="PNG")
        preview_images.append(preview_buffer.getvalue())
        page_text = ocr_image_to_text(enhanced_image)
        if page_text:
            extracted_pages.append(page_text)
            page_records.append({"page": page_number, "text": page_text})

    if not page_records:
        raise RuntimeError("OCR could not extract readable text from the selected scans.")

    base_name = clean_filename(requested_name or f"scan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf")
    if not base_name.lower().endswith(".pdf"):
        base_name = f"{base_name}.pdf"
    pdf_bytes = images_to_pdf_bytes(enhanced_images)
    return InMemoryUpload(base_name, pdf_bytes), "\n\n".join(extracted_pages), len(enhanced_images), page_records, preview_images


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []

    chunks: list[str] = []
    start = 0
    text_length = len(normalized)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = normalized[start:end]
        if end < text_length:
            sentence_break = max(chunk.rfind(". "), chunk.rfind("? "), chunk.rfind("! "))
            if sentence_break > chunk_size * 0.6:
                end = start + sentence_break + 1
                chunk = normalized[start:end]
        chunks.append(chunk.strip())
        if end >= text_length:
            break
        start = max(end - overlap, start + 1)
    return [chunk for chunk in chunks if chunk]


def chunk_page_records(page_records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    chunk_records: list[dict[str, Any]] = []
    for page_record in page_records:
        page_number = page_record.get("page")
        for chunk in chunk_text(page_record.get("text", "")):
            chunk_records.append({"text": chunk, "page": page_number})
    return chunk_records


def normalize_chunk_record(raw_chunk: Any) -> dict[str, Any]:
    if isinstance(raw_chunk, dict):
        return {
            "text": str(raw_chunk.get("text", "")).strip(),
            "page": raw_chunk.get("page"),
        }
    return {"text": str(raw_chunk).strip(), "page": None}


def ui_text(language: str, key: str) -> str:
    language_code = LANGUAGE_CONFIG.get(language, LANGUAGE_CONFIG["English"])["code"]
    bundle = TEXTS.get(key, {})
    return bundle.get(language_code) or bundle.get("en") or key


def page_label(page_number: int | None) -> str:
    return f"Page {page_number}" if page_number else "Page unknown"


class LocalAssistantService:
    def __init__(self) -> None:
        self.provider = LOCAL_PROVIDER

    def generate_answer(self, question: str, contexts: list[str]) -> str:
        if not contexts:
            return INFO_NOT_FOUND
        unique_contexts = list(dict.fromkeys(chunk.strip() for chunk in contexts if chunk.strip()))
        return "\n\n".join(unique_contexts[:3]) if unique_contexts else INFO_NOT_FOUND

    def transcribe_audio(self, audio_bytes: bytes) -> str:
        if sr is None:
            raise RuntimeError("Offline voice transcription requires `SpeechRecognition` and `pocketsphinx`.")
        recognizer = sr.Recognizer()
        audio_file = io.BytesIO(audio_bytes)
        try:
            with sr.AudioFile(audio_file) as source:
                audio_data = recognizer.record(source)
            return recognizer.recognize_sphinx(audio_data).strip()
        except sr.UnknownValueError:
            return ""
        except Exception as exc:
            raise RuntimeError(f"Offline voice transcription failed: {exc}") from exc


def get_available_providers() -> list[str]:
    return [LOCAL_PROVIDER]


def resolve_provider(preferred_provider: str | None, available_providers: list[str]) -> str:
    return LOCAL_PROVIDER


def get_ai_service(provider: str) -> LocalAssistantService:
    return LocalAssistantService()


def save_chunk_assets(user_id: int, doc_token: str, chunks: list[dict[str, Any]]) -> Path:
    user_vector_dir = VECTOR_DIR / f"user_{user_id}"
    user_vector_dir.mkdir(parents=True, exist_ok=True)

    chunks_path = user_vector_dir / f"{doc_token}_chunks.json"
    chunks_path.write_text(json.dumps(chunks, ensure_ascii=True), encoding="utf-8")
    return chunks_path


def sync_document_to_cloud(user_id: int, saved_name: str, source_path: Path) -> str:
    user_cloud_dir = CLOUD_DIR / f"user_{user_id}"
    user_cloud_dir.mkdir(parents=True, exist_ok=True)
    cloud_path = user_cloud_dir / saved_name
    shutil.copy2(source_path, cloud_path)
    return str(cloud_path)


def get_user_cloud_file_count(user_id: int) -> int:
    user_cloud_dir = CLOUD_DIR / f"user_{user_id}"
    if not user_cloud_dir.exists():
        return 0
    return len(list(user_cloud_dir.glob("*.pdf")))


def store_document_bytes(
    user_id: int,
    uploaded_file: Any,
    ai_service: LocalAssistantService,
    file_bytes: bytes,
    extracted_text: str,
    page_count: int,
    page_records: list[dict[str, Any]],
) -> tuple[bool, str]:
    if not extracted_text.strip():
        return False, f"No extractable text found in {uploaded_file.name}."

    chunks = chunk_page_records(page_records)
    if not chunks:
        return False, f"Unable to create searchable chunks for {uploaded_file.name}."

    file_hash = compute_file_hash(file_bytes)

    user_upload_dir = UPLOAD_DIR / f"user_{user_id}"
    user_upload_dir.mkdir(parents=True, exist_ok=True)

    doc_token = uuid.uuid4().hex
    safe_name = clean_filename(uploaded_file.name)
    saved_name = f"{doc_token}_{safe_name}"
    file_path = user_upload_dir / saved_name
    chunks_path = save_chunk_assets(user_id, doc_token, chunks)
    file_path.write_bytes(file_bytes)
    cloud_path = sync_document_to_cloud(user_id, saved_name, file_path)

    with db_session() as session:
        duplicate = session.scalar(
            select(Document).where(
                Document.user_id == user_id,
                Document.file_hash == file_hash,
                Document.original_name == uploaded_file.name,
            )
        )
        if duplicate:
            return False, f"{uploaded_file.name} was already uploaded in your workspace."

        session.add(
            Document(
                user_id=user_id,
                original_name=uploaded_file.name,
                saved_name=saved_name,
                file_path=str(file_path),
                extracted_text=extracted_text,
                chunks_path=str(chunks_path),
                embeddings_path="",
                index_path="",
                chunk_count=len(chunks),
                created_at=utc_now(),
                file_hash=file_hash,
                page_count=page_count,
                embedding_provider=ai_service.provider,
                embedding_model="",
            )
        )

    return True, f"Processed {uploaded_file.name} with {len(chunks)} searchable chunks. Cloud copy saved to {cloud_path}."


def save_document(user_id: int, uploaded_file: Any, ai_service: LocalAssistantService) -> tuple[bool, str]:
    file_bytes = uploaded_file.getvalue()
    extracted_text, page_count, page_records = extract_page_records(file_bytes)
    return store_document_bytes(user_id, uploaded_file, ai_service, file_bytes, extracted_text, page_count, page_records)


def load_document_chunks(document: dict[str, Any]) -> list[dict[str, Any]]:
    chunks_path = Path(document["chunks_path"])
    if not chunks_path.exists():
        return []
    return [normalize_chunk_record(chunk) for chunk in json.loads(chunks_path.read_text(encoding="utf-8"))]


def rebuild_document_chunks(document_id: int, ai_service: LocalAssistantService) -> dict[str, Any]:
    with db_session() as session:
        document = session.get(Document, document_id)
        if not document:
            raise RuntimeError("Document not found.")

        if document.file_path and Path(document.file_path).exists():
            extracted_text, _, page_records = extract_page_records(Path(document.file_path).read_bytes())
            chunks = chunk_page_records(page_records)
            if extracted_text.strip():
                document.extracted_text = extracted_text
        else:
            chunks = [{"text": chunk, "page": None} for chunk in chunk_text(document.extracted_text)]
        if not chunks:
            raise RuntimeError(f"Unable to chunk {document.original_name} for retrieval.")

        token = Path(document.saved_name).stem if document.saved_name else uuid.uuid4().hex
        chunks_path = save_chunk_assets(document.user_id, token, chunks)
        document.chunks_path = str(chunks_path)
        document.embeddings_path = ""
        document.index_path = ""
        document.chunk_count = len(chunks)
        document.embedding_provider = ai_service.provider
        document.embedding_model = ""
        session.add(document)
        session.flush()
        session.refresh(document)
        return serialize_document(document)


def ensure_documents_ready(documents: list[dict[str, Any]], ai_service: LocalAssistantService) -> list[dict[str, Any]]:
    aligned: list[dict[str, Any]] = []
    for document in documents:
        chunks_exist = bool(document.get("chunks_path")) and Path(document["chunks_path"]).exists()
        if not chunks_exist:
            aligned.append(rebuild_document_chunks(document["id"], ai_service))
        else:
            aligned.append(document)
    return aligned


def tokenize_text(text_value: str) -> list[str]:
    return [
        token
        for token in re.findall(r"[A-Za-z0-9']+", text_value.lower())
        if token and token not in STOP_WORDS and len(token) > 1
    ]


def score_candidate(question_tokens: list[str], candidate: str) -> int:
    if not candidate.strip():
        return 0
    candidate_tokens = set(tokenize_text(candidate))
    if not candidate_tokens:
        return 0
    overlap = sum(1 for token in question_tokens if token in candidate_tokens)
    phrase_bonus = 2 if len(question_tokens) >= 2 and " ".join(question_tokens[:2]) in candidate.lower() else 0
    digit_bonus = sum(1 for token in question_tokens if token.isdigit() and token in candidate)
    return overlap + phrase_bonus + digit_bonus


def best_excerpt_for_chunk(question: str, chunk: str) -> tuple[str, int]:
    question_tokens = tokenize_text(question)
    sentences = [value.strip() for value in re.split(r"(?<=[.!?])\s+|\n+", chunk) if value.strip()]
    candidates = sentences or [chunk.strip()]
    best_text = ""
    best_score = 0
    for candidate in candidates:
        score = score_candidate(question_tokens, candidate)
        if score > best_score:
            best_text = candidate
            best_score = score
    if best_score == 0 and not question_tokens and candidates:
        return textwrap.shorten(candidates[0], width=420, placeholder="..."), 1
    excerpt = best_text or chunk.strip()
    return textwrap.shorten(excerpt, width=420, placeholder="..."), best_score


def retrieve_relevant_context(
    ai_service: LocalAssistantService,
    selected_docs: list[dict[str, Any]],
    question: str,
) -> tuple[list[str], list[str], float, list[dict[str, Any]]]:
    if not selected_docs:
        return [], [], 0.0, []

    ranked_matches: list[dict[str, Any]] = []

    for document in selected_docs:
        chunks = load_document_chunks(document)
        for chunk in chunks:
            excerpt, score = best_excerpt_for_chunk(question, chunk["text"])
            if score >= MIN_MATCH_SCORE:
                ranked_matches.append(
                    {
                        "score": score,
                        "source": document["original_name"],
                        "page": chunk.get("page"),
                        "excerpt": excerpt,
                    }
                )

    ranked_matches.sort(key=lambda item: item["score"], reverse=True)
    top_matches = ranked_matches[:MAX_CONTEXT_CHUNKS]
    contexts = [match["excerpt"] for match in top_matches]
    sources = sorted({f'{match["source"]} ({page_label(match["page"])})' for match in top_matches})
    best_score = float(top_matches[0]["score"]) if top_matches else 0.0
    return contexts, sources, best_score, top_matches


def build_chat_pdf(messages: list[dict[str, Any]], title: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    _, height = A4
    y = height - 50

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, title)
    y -= 30
    pdf.setFont("Helvetica", 10)

    for message in messages:
        speaker = "User" if message["role"] == "user" else "AI"
        lines = textwrap.wrap(f"{speaker}: {message['content']}", width=95) or [f"{speaker}:"]
        for line in lines:
            if y < 50:
                pdf.showPage()
                pdf.setFont("Helvetica", 10)
                y = height - 50
            pdf.drawString(50, y, line)
            y -= 14
        y -= 8

    pdf.save()
    buffer.seek(0)
    return buffer.read()


def build_chat_txt(messages: list[dict[str, Any]], title: str) -> bytes:
    lines = [title, "=" * len(title), ""]
    for message in messages:
        speaker = "User" if message["role"] == "user" else "Assistant"
        lines.append(f"{speaker}: {message['content']}")
        metadata = message.get("metadata") or {}
        sources = metadata.get("sources") or []
        if sources:
            lines.append(f"Sources: {', '.join(sources)}")
        lines.append("")
    return "\n".join(lines).encode("utf-8")


def build_chat_docx(messages: list[dict[str, Any]], title: str) -> bytes:
    if DocxFile is None:
        raise RuntimeError("DOCX export requires `python-docx`.")
    document = DocxFile()
    document.add_heading(title, level=1)
    for message in messages:
        speaker = "User" if message["role"] == "user" else "Assistant"
        paragraph = document.add_paragraph()
        paragraph.add_run(f"{speaker}: ").bold = True
        paragraph.add_run(message["content"])
        metadata = message.get("metadata") or {}
        sources = metadata.get("sources") or []
        if sources:
            document.add_paragraph(f"Sources: {', '.join(sources)}")
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def split_sentences(text_value: str) -> list[str]:
    return [segment.strip() for segment in re.split(r"(?<=[.!?])\s+|\n+", text_value) if segment.strip()]


def gather_document_sentences(selected_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    sentence_records: list[dict[str, Any]] = []
    for document in selected_docs:
        for chunk in load_document_chunks(document):
            for sentence in split_sentences(chunk["text"]):
                if len(sentence.split()) >= 6:
                    sentence_records.append(
                        {
                            "source": document["original_name"],
                            "page": chunk.get("page"),
                            "text": sentence,
                        }
                    )
    return sentence_records


def rank_sentence_records(sentence_records: list[dict[str, Any]]) -> list[dict[str, Any]]:
    token_counter: Counter[str] = Counter()
    for record in sentence_records:
        token_counter.update(tokenize_text(record["text"]))

    ranked: list[dict[str, Any]] = []
    seen: set[str] = set()
    for record in sentence_records:
        sentence = record["text"]
        sentence_key = sentence.lower()
        if sentence_key in seen:
            continue
        seen.add(sentence_key)
        tokens = tokenize_text(sentence)
        score = sum(token_counter[token] for token in tokens) + min(len(tokens), 18)
        ranked.append({**record, "score": score})
    ranked.sort(key=lambda item: item["score"], reverse=True)
    return ranked


def build_summary(selected_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return rank_sentence_records(gather_document_sentences(selected_docs))[:5]


def build_notes(selected_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return rank_sentence_records(gather_document_sentences(selected_docs))[:MAX_STUDY_SENTENCES]


def extract_subject_and_fact(sentence: str) -> tuple[str, str]:
    patterns = (" is ", " are ", " refers to ", " means ", " includes ", " uses ")
    for pattern in patterns:
        if pattern in sentence:
            left, right = sentence.split(pattern, 1)
            return left.strip(" .:-"), right.strip(" .")
    words = sentence.split()
    subject = " ".join(words[:4]).strip(" .:-")
    fact = " ".join(words[4:]).strip(" .")
    return subject, fact


def build_important_questions(selected_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    questions: list[dict[str, Any]] = []
    for record in rank_sentence_records(gather_document_sentences(selected_docs))[:6]:
        subject, _ = extract_subject_and_fact(record["text"])
        prompt = f"What does the document say about {subject}?" if subject else "What is an important idea from the document?"
        questions.append({**record, "question": prompt})
    return questions


def build_mcqs(selected_docs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    ranked = rank_sentence_records(gather_document_sentences(selected_docs))
    distractor_pool = [record["text"] for record in ranked[4:12]]
    mcqs: list[dict[str, Any]] = []
    for index, record in enumerate(ranked[:4], start=1):
        subject, fact = extract_subject_and_fact(record["text"])
        correct_option = textwrap.shorten(fact or record["text"], width=110, placeholder="...")
        distractors = [textwrap.shorten(option, width=110, placeholder="...") for option in distractor_pool[:3]]
        while len(distractors) < 3:
            distractors.append("Not stated in this excerpt")
        options = [correct_option, *distractors[:3]]
        prompt = f"What best matches {subject}?" if subject else f"Which statement is supported by the PDF? ({index})"
        mcqs.append(
            {
                **record,
                "question": prompt,
                "options": options,
                "answer": "A",
            }
        )
    return mcqs


def format_study_lines(records: list[dict[str, Any]], include_question: bool = False) -> list[str]:
    lines: list[str] = []
    for index, record in enumerate(records, start=1):
        prefix = record.get("question") if include_question else record["text"]
        lines.append(f"{index}. {prefix}")
        lines.append(f"Source: {record['source']} - {page_label(record.get('page'))}")
    return lines


def get_user_analytics(user_id: int) -> dict[str, int]:
    documents = get_user_documents(user_id)
    chats = get_user_chats(user_id)
    question_count = 0
    for chat in chats:
        question_count += sum(1 for message in get_messages(chat["id"]) if message["role"] == "user")
    return {
        "pdf_count": len(documents),
        "chat_count": len(chats),
        "question_count": question_count,
        "page_count": sum(int(doc["page_count"]) for doc in documents),
        "chunk_count": sum(int(doc["chunk_count"]) for doc in documents),
        "cloud_count": get_user_cloud_file_count(user_id),
    }


def render_speech_player(text_value: str, speech_locale: str, language: str) -> None:
    safe_text = json.dumps(text_value)
    safe_locale = json.dumps(speech_locale)
    safe_labels = json.dumps(
        {
            "play": "Play answer",
            "pause": "Pause",
            "resume": "Resume",
            "stop": "Stop",
            "status": "Status",
            "ready": "Ready",
            "speaking": "Speaking",
            "paused": "Paused",
            "stopped": "Stopped",
            "finished": "Finished",
            "ready_message": "Latest answer is ready for playback.",
            "speaking_message": "Reading the latest AI answer aloud.",
            "paused_message": "Playback paused. Resume anytime.",
            "stopped_message": "Playback stopped.",
            "finished_message": "Playback finished.",
            "error_message": "Playback could not continue.",
            "unsupported_message": "Speech playback is not supported in this browser.",
            "empty_message": ui_text(language, "voice_output_empty"),
        }
    )
    components.html(
        f"""
        <div class="voice-player">
            <div class="voice-player__status">
                <span class="voice-player__status-label">Status</span>
                <span id="voice-status-pill" class="voice-player__pill">Ready</span>
            </div>
            <div class="voice-player__controls">
                <button id="voice-play" class="voice-player__button voice-player__button--primary" type="button">▶ Play answer</button>
                <button id="voice-pause" class="voice-player__button" type="button">⏸ Pause</button>
                <button id="voice-stop" class="voice-player__button voice-player__button--danger" type="button">⏹ Stop</button>
            </div>
            <div id="voice-helper" class="voice-player__helper"></div>
        </div>
        <style>
        body {{
            margin: 0;
            font-family: "Segoe UI", Arial, sans-serif;
            background: transparent;
            color: #eaf2ff;
        }}
        .voice-player {{
            border: 1px solid rgba(255,255,255,0.10);
            background: linear-gradient(145deg, rgba(14,26,44,0.95), rgba(10,18,31,0.92));
            border-radius: 18px;
            padding: 14px;
            box-sizing: border-box;
        }}
        .voice-player__status {{
            display: flex;
            align-items: center;
            justify-content: space-between;
            gap: 12px;
            margin-bottom: 12px;
        }}
        .voice-player__status-label {{
            font-size: 0.82rem;
            letter-spacing: 0.02em;
            color: #9db0cb;
        }}
        .voice-player__pill {{
            padding: 6px 12px;
            border-radius: 999px;
            background: rgba(61, 217, 176, 0.14);
            color: #7ef0d2;
            font-size: 0.8rem;
            font-weight: 600;
        }}
        .voice-player__controls {{
            display: grid;
            grid-template-columns: repeat(3, minmax(0, 1fr));
            gap: 10px;
        }}
        .voice-player__button {{
            border: 1px solid rgba(255,255,255,0.12);
            background: rgba(255,255,255,0.05);
            color: #edf4ff;
            border-radius: 12px;
            padding: 11px 10px;
            font-size: 0.92rem;
            font-weight: 600;
            cursor: pointer;
            transition: transform 0.15s ease, background 0.15s ease, border-color 0.15s ease;
        }}
        .voice-player__button:hover:not(:disabled) {{
            transform: translateY(-1px);
            background: rgba(255,255,255,0.09);
            border-color: rgba(255,255,255,0.18);
        }}
        .voice-player__button:disabled {{
            opacity: 0.5;
            cursor: not-allowed;
        }}
        .voice-player__button--primary {{
            background: linear-gradient(135deg, rgba(61, 217, 176, 0.92), rgba(79, 140, 255, 0.88));
            color: #08111e;
            border-color: transparent;
        }}
        .voice-player__button--danger {{
            background: rgba(255, 107, 107, 0.13);
            color: #ffb2b2;
        }}
        .voice-player__helper {{
            min-height: 18px;
            margin-top: 10px;
            color: #98abc7;
            font-size: 0.8rem;
        }}
        @media (max-width: 640px) {{
            .voice-player__controls {{
                grid-template-columns: 1fr;
            }}
        }}
        </style>
        <script>
        const text = {safe_text};
        const locale = {safe_locale};
        const labels = {safe_labels};
        const synth = window.speechSynthesis;
        const statusLabel = document.querySelector(".voice-player__status-label");
        const statusPill = document.getElementById("voice-status-pill");
        const helper = document.getElementById("voice-helper");
        const playButton = document.getElementById("voice-play");
        const pauseButton = document.getElementById("voice-pause");
        const stopButton = document.getElementById("voice-stop");

        function setStatus(statusName, helperText = "") {{
            statusLabel.textContent = labels.status;
            statusPill.textContent = labels[statusName] || statusName;
            helper.textContent = helperText;
        }}

        function updateButtons() {{
            const hasText = Boolean(text);
            const isSpeaking = Boolean(synth && synth.speaking);
            const isPaused = Boolean(synth && synth.paused);
            playButton.disabled = !hasText || isSpeaking;
            pauseButton.disabled = !hasText || !isSpeaking;
            pauseButton.textContent = isPaused ? `⏵ ${{labels.resume}}` : `⏸ ${{labels.pause}}`;
            stopButton.disabled = !hasText || (!isSpeaking && !isPaused);
        }}

        function stopSpeech(statusName = "stopped", helperText = labels.stopped_message) {{
            if (synth) {{
                synth.cancel();
            }}
            setStatus(statusName, helperText);
            updateButtons();
        }}

        function speakText() {{
            if (!text) {{
                setStatus("ready", labels.empty_message);
                updateButtons();
                return;
            }}
            if (!synth) {{
                setStatus("stopped", labels.unsupported_message);
                playButton.disabled = true;
                pauseButton.disabled = true;
                stopButton.disabled = true;
                return;
            }}

            synth.cancel();
            const utterance = new SpeechSynthesisUtterance(text);
            utterance.rate = 1;
            utterance.pitch = 1;
            utterance.lang = locale;
            utterance.onstart = () => {{
                setStatus("speaking", labels.speaking_message);
                updateButtons();
            }};
            utterance.onpause = () => {{
                setStatus("paused", labels.paused_message);
                updateButtons();
            }};
            utterance.onresume = () => {{
                setStatus("speaking", labels.speaking_message);
                updateButtons();
            }};
            utterance.onend = () => {{
                setStatus("finished", labels.finished_message);
                updateButtons();
            }};
            utterance.onerror = () => {{
                setStatus("stopped", labels.error_message);
                updateButtons();
            }};
            synth.speak(utterance);
            setStatus("speaking", labels.speaking_message);
            updateButtons();
        }}

        playButton.addEventListener("click", speakText);
        pauseButton.addEventListener("click", () => {{
            if (!synth || !synth.speaking) {{
                updateButtons();
                return;
            }}
            if (synth.paused) {{
                synth.resume();
                setStatus("speaking", labels.speaking_message);
            }} else {{
                synth.pause();
                setStatus("paused", labels.paused_message);
            }}
            updateButtons();
        }});
        stopButton.addEventListener("click", () => stopSpeech());

        window.addEventListener("beforeunload", () => {{
            if (synth) {{
                synth.cancel();
            }}
        }});

        if (text) {{
            setStatus("ready", labels.ready_message);
        }} else {{
            setStatus("ready", labels.empty_message);
        }}
        playButton.textContent = `▶ ${{labels.play}}`;
        stopButton.textContent = `⏹ ${{labels.stop}}`;
        updateButtons();
        </script>
        """,
        height=170,
    )


def inject_styles(theme: str) -> None:
    st.set_page_config(page_title=APP_TITLE, page_icon="📚", layout="wide")
    palettes = {
        "dark": {
            "bg": "#08111e",
            "bg2": "#0d1728",
            "panel": "rgba(12, 22, 38, 0.88)",
            "panel_soft": "rgba(17, 31, 53, 0.82)",
            "text": "#edf4ff",
            "muted": "#98abc7",
            "border": "rgba(255,255,255,0.09)",
            "accent": "#3dd9b0",
            "accent2": "#4f8cff",
            "accent_glow": "rgba(61, 217, 176, 0.22)",
            "shadow": "0 24px 80px rgba(0,0,0,0.28)",
            "chat_user": "rgba(79, 140, 255, 0.15)",
            "chat_ai": "rgba(255, 255, 255, 0.04)",
        },
        "light": {
            "bg": "#f4f7fb",
            "bg2": "#eef4ff",
            "panel": "rgba(255, 255, 255, 0.88)",
            "panel_soft": "rgba(247, 250, 255, 0.96)",
            "text": "#162033",
            "muted": "#5e6f89",
            "border": "rgba(16, 24, 40, 0.08)",
            "accent": "#0f9d7a",
            "accent2": "#2c72ff",
            "accent_glow": "rgba(44, 114, 255, 0.16)",
            "shadow": "0 18px 60px rgba(35,52,82,0.12)",
            "chat_user": "rgba(44, 114, 255, 0.10)",
            "chat_ai": "rgba(22, 32, 51, 0.03)",
        },
    }
    colors = palettes.get(theme, palettes["dark"])
    st.markdown(
        f"""
        <style>
        .stApp {{
            background:
                radial-gradient(circle at top left, {colors["accent_glow"]}, transparent 24%),
                radial-gradient(circle at top right, rgba(255, 146, 88, 0.10), transparent 20%),
                linear-gradient(180deg, {colors["bg2"]} 0%, {colors["bg"]} 100%);
            color: {colors["text"]};
        }}
        [data-testid="stSidebar"] {{
            background: {colors["panel"]};
            border-right: 1px solid {colors["border"]};
        }}
        .block-container {{
            max-width: 1360px;
            padding-top: 1.6rem;
            padding-bottom: 2rem;
        }}
        h1, h2, h3, h4, h5, h6, p, span, label, div {{
            color: {colors["text"]};
        }}
        .hero-card, .panel-card, .metric-card, .chat-shell {{
            background: {colors["panel"]};
            border: 1px solid {colors["border"]};
            border-radius: 24px;
            box-shadow: {colors["shadow"]};
            backdrop-filter: blur(18px);
        }}
        .hero-card {{
            padding: 1.45rem;
            margin-bottom: 1rem;
        }}
        .panel-card {{
            padding: 1rem 1.1rem;
        }}
        .metric-card {{
            padding: 1rem 1.1rem;
            height: 100%;
        }}
        .hero-title {{
            font-size: 2.55rem;
            line-height: 1;
            font-weight: 800;
            letter-spacing: -0.05em;
            margin-bottom: 0.35rem;
        }}
        .hero-subtitle, .helper-text {{
            color: {colors["muted"]};
        }}
        .pill {{
            display: inline-block;
            padding: 0.4rem 0.78rem;
            margin: 0.22rem 0.45rem 0.22rem 0;
            border-radius: 999px;
            border: 1px solid {colors["border"]};
            background: {colors["panel_soft"]};
            color: {colors["text"]};
            font-size: 0.92rem;
        }}
        .metric-value {{
            font-size: 1.6rem;
            font-weight: 800;
            margin-top: 0.2rem;
        }}
        .chat-shell {{
            padding: 1rem;
            margin-bottom: 1rem;
        }}
        .stChatMessage {{
            border-radius: 20px;
            border: 1px solid {colors["border"]};
            padding: 0.35rem;
            background: {colors["chat_ai"]};
        }}
        .stChatMessage[data-testid="chatAvatarIcon-user"] {{
            background: transparent;
        }}
        [data-testid="stChatMessageContent"] {{
            color: {colors["text"]};
        }}
        .user-bubble {{
            background: {colors["chat_user"]};
            border: 1px solid {colors["border"]};
            border-radius: 18px;
            padding: 0.7rem 0.9rem;
        }}
        .assistant-bubble {{
            background: {colors["chat_ai"]};
            border: 1px solid {colors["border"]};
            border-radius: 18px;
            padding: 0.7rem 0.9rem;
        }}
        .stTextInput input, .stTextArea textarea, .stSelectbox div[data-baseweb="select"] > div,
        .stMultiSelect div[data-baseweb="select"] > div, .stFileUploader, .stAudioInput {{
            border-radius: 16px !important;
            background: {colors["panel_soft"]} !important;
            border: 1px solid {colors["border"]} !important;
            color: {colors["text"]} !important;
        }}
        .stButton > button, .stDownloadButton > button {{
            border-radius: 999px;
            border: none;
            background: linear-gradient(90deg, {colors["accent"]}, {colors["accent2"]});
            color: white;
            font-weight: 700;
            box-shadow: 0 12px 28px {colors["accent_glow"]};
        }}
        .stButton > button:hover, .stDownloadButton > button:hover {{
            filter: brightness(1.03);
        }}
        .stCaption, .stMarkdown small {{
            color: {colors["muted"]};
        }}
        @media (max-width: 920px) {{
            .hero-title {{
                font-size: 2rem;
            }}
            .block-container {{
                padding-top: 1rem;
            }}
        }}
        </style>
        """,
        unsafe_allow_html=True,
    )


def init_session_state() -> None:
    st.session_state.setdefault("user", None)
    st.session_state.setdefault("active_chat_id", None)
    st.session_state.setdefault("theme", "dark")
    st.session_state.setdefault("pending_prompt", "")
    st.session_state.setdefault("last_answer_to_speak", "")
    st.session_state.setdefault("language", "English")
    st.session_state.setdefault("page_search_query", "")
    st.session_state.setdefault("confirm_clear_chat_id", None)
    st.session_state.setdefault("confirm_clear_all", False)
    st.session_state.setdefault("camera_scan_pages", [])
    st.session_state.setdefault("camera_scan_hashes", [])
    st.session_state.setdefault("camera_scan_enabled", False)
    st.session_state.setdefault("camera_scan_input_key", 0)


def reset_chat_clear_state() -> None:
    st.session_state.confirm_clear_chat_id = None
    st.session_state.confirm_clear_all = False


def reset_scan_state() -> None:
    st.session_state.camera_scan_pages = []
    st.session_state.camera_scan_hashes = []
    st.session_state.camera_scan_enabled = False
    st.session_state.camera_scan_input_key = int(st.session_state.get("camera_scan_input_key", 0)) + 1


def sync_user_session(user: dict[str, Any]) -> None:
    st.session_state.user = user
    st.session_state.theme = user.get("theme_preference", "dark")
    reset_scan_state()


def render_auth_screen() -> None:
    left_col, right_col = st.columns([1.15, 0.85], gap="large")
    with left_col:
        st.markdown(
            """
            <div class="hero-card">
                <div class="hero-title">Private PDF intelligence for every user</div>
                <div class="hero-subtitle">
                    Upload one or many PDFs, search them locally offline, and chat through a clean SaaS-style workspace
                    that only answers from your own documents.
                </div>
                <div style="margin-top: 1rem;">
                    <span class="pill">Login and signup</span>
                    <span class="pill">Multi-PDF upload</span>
                    <span class="pill">Offline PDF search</span>
                    <span class="pill">Voice questions</span>
                    <span class="pill">Saved chat history</span>
                    <span class="pill">Deployment ready</span>
                </div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        metrics = st.columns(3)
        metric_data = [
            ("Document Search", "Offline"),
            ("Database", "Postgres Ready"),
            ("Answer Policy", "Strict PDF Only"),
        ]
        for column, (label, value) in zip(metrics, metric_data):
            with column:
                st.markdown(
                    f"""
                    <div class="metric-card">
                        <div class="helper-text">{label}</div>
                        <div class="metric-value">{value}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

    with right_col:
        st.markdown('<div class="panel-card">', unsafe_allow_html=True)
        st.subheader("Welcome back")
        mode = st.radio(
            "Authentication",
            ["login", "signup"],
            horizontal=True,
            label_visibility="collapsed",
            format_func=lambda value: "Sign in" if value == "login" else "Create account",
        )
        with st.form("auth_form", clear_on_submit=False):
            email = st.text_input("Email")
            password = st.text_input("Password", type="password")
            confirm_password = st.text_input("Confirm password", type="password") if mode == "signup" else ""
            submitted = st.form_submit_button("Create account" if mode == "signup" else "Sign in", use_container_width=True)

        if submitted:
            if mode == "signup":
                if password != confirm_password:
                    st.error("Passwords do not match.")
                else:
                    ok, message = create_user(email, password)
                    st.success(message) if ok else st.error(message)
            else:
                user = authenticate_user(email, password)
                if user:
                    sync_user_session(user)
                    st.session_state.active_chat_id = None
                    st.rerun()
                else:
                    st.error("Invalid email or password.")
        st.markdown("</div>", unsafe_allow_html=True)


def render_sidebar(user: dict[str, Any]) -> tuple[list[dict[str, Any]], list[int], str, str]:
    with st.sidebar:
        st.markdown(f"## {APP_TITLE}")
        st.caption(f"Signed in as `{user['email']}`")

        theme = st.radio(
            "Theme",
            options=["dark", "light"],
            index=0 if st.session_state.theme == "dark" else 1,
            horizontal=True,
        )
        if theme != st.session_state.theme:
            st.session_state.theme = theme
            updated = update_user_preferences(user["id"], theme=theme)
            st.session_state.user = updated
            st.rerun()

        language = st.selectbox(
            ui_text(st.session_state.language, "language"),
            options=list(LANGUAGE_CONFIG.keys()),
            index=list(LANGUAGE_CONFIG.keys()).index(st.session_state.language),
        )
        if language != st.session_state.language:
            st.session_state.language = language
            st.rerun()

        available_providers = get_available_providers()
        provider = resolve_provider(user.get("preferred_provider", DEFAULT_PROVIDER), available_providers)
        st.caption(ui_text(language, "answer_mode"))
        st.caption(ui_text(language, "local_run"))
        st.caption(f"{ui_text(language, 'cloud')}: `{CLOUD_DIR}`")
        st.caption(f"{get_user_cloud_file_count(user['id'])} mirrored PDFs")

        if provider != user.get("preferred_provider", DEFAULT_PROVIDER):
            updated = update_user_preferences(user["id"], provider=provider)
            st.session_state.user = updated
            user = updated
            st.rerun()

        documents = get_user_documents(user["id"])
        chats = get_user_chats(user["id"])

        st.markdown("### Documents")
        selected_doc_ids: list[int] = []
        active_chat = get_chat(st.session_state.active_chat_id, user["id"]) if st.session_state.active_chat_id else None
        default_ids = json.loads(active_chat["selected_document_ids"]) if active_chat else []
        if documents:
            options = {f"{doc['original_name']} - {doc['chunk_count']} chunks": doc["id"] for doc in documents}
            labels = [label for label, document_id in options.items() if document_id in default_ids]
            selected_labels = st.multiselect(
                "Choose PDFs for this chat",
                list(options.keys()),
                default=labels,
                help="Answers will be restricted to the selected PDFs only.",
            )
            selected_doc_ids = [options[label] for label in selected_labels]
            if active_chat and selected_doc_ids != default_ids:
                update_chat_documents(active_chat["id"], selected_doc_ids)
        else:
            st.info("Upload at least one PDF to start chatting.")

        uploads = st.file_uploader("Upload one or more PDFs", type="pdf", accept_multiple_files=True)
        if st.button("Process uploads", use_container_width=True):
            if not uploads:
                st.warning("Choose at least one PDF.")
            else:
                try:
                    ai_service = get_ai_service(provider)
                    processed = 0
                    for uploaded_file in uploads:
                        ok, message = save_document(user["id"], uploaded_file, ai_service)
                        st.success(message) if ok else st.warning(message)
                        processed += int(ok)
                    if processed:
                        st.rerun()
                except Exception as exc:
                    st.error(str(exc))

        render_scan_import_tools(user["id"], provider)

        st.markdown("### Chat History")
        if st.button("New chat", use_container_width=True):
            st.session_state.active_chat_id = None
            reset_chat_clear_state()
            st.rerun()

        with st.expander("Manage history", expanded=False):
            if active_chat:
                armed_for_current = st.session_state.confirm_clear_chat_id == active_chat["id"]
                current_label = f"Confirm delete: {active_chat['title']}" if armed_for_current else f"🗑️ {ui_text(language, 'clear_current_chat')}"
                if st.button(current_label, key="clear_current_chat", use_container_width=True):
                    if armed_for_current:
                        delete_chat(active_chat["id"], user["id"])
                        st.session_state.active_chat_id = None
                        st.session_state.pending_prompt = ""
                        st.session_state.last_answer_to_speak = ""
                        reset_chat_clear_state()
                        st.rerun()
                    else:
                        st.session_state.confirm_clear_chat_id = active_chat["id"]
                        st.session_state.confirm_clear_all = False
                        st.rerun()
            else:
                st.button(
                    f"🗑️ {ui_text(language, 'clear_current_chat')}",
                    key="clear_current_chat_disabled",
                    use_container_width=True,
                    disabled=True,
                )

            armed_for_all = st.session_state.confirm_clear_all
            all_label = "Confirm delete all chats" if armed_for_all else f"🗑️ {ui_text(language, 'clear_all_chats')}"
            if st.button(all_label, key="clear_all_chats", use_container_width=True, disabled=not chats):
                if armed_for_all:
                    delete_all_chats(user["id"])
                    st.session_state.active_chat_id = None
                    st.session_state.pending_prompt = ""
                    st.session_state.last_answer_to_speak = ""
                    reset_chat_clear_state()
                    st.rerun()
                else:
                    st.session_state.confirm_clear_all = True
                    st.session_state.confirm_clear_chat_id = None
                    st.rerun()

            if st.session_state.confirm_clear_chat_id or st.session_state.confirm_clear_all:
                if st.button("Cancel", key="cancel_clear_history", use_container_width=True):
                    reset_chat_clear_state()
                    st.rerun()

        for chat in chats:
            label = f"{chat['title']}"
            if st.button(label, key=f"chat_{chat['id']}", use_container_width=True):
                st.session_state.active_chat_id = chat["id"]
                reset_chat_clear_state()
                st.rerun()

        if st.button("Sign out", use_container_width=True):
            st.session_state.user = None
            st.session_state.active_chat_id = None
            st.session_state.pending_prompt = ""
            reset_scan_state()
            reset_chat_clear_state()
            st.rerun()

    return documents, selected_doc_ids, provider, language


def render_workspace_header(documents: list[dict[str, Any]], chats: list[dict[str, Any]], provider: str, language: str, analytics: dict[str, int]) -> None:
    metrics = st.columns(3)
    values = [
        ("Indexed PDFs", str(len(documents))),
        ("Saved chats", str(analytics["chat_count"])),
        ("Questions asked", str(analytics["question_count"])),
    ]
    for column, (label, value) in zip(metrics, values):
        with column:
            st.markdown(
                f"""
                <div class="metric-card">
                    <div class="helper-text">{label}</div>
                    <div class="metric-value">{value}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )

    st.markdown(
        f"""
        <div class="hero-card">
            <div class="hero-title">Chat with your PDFs</div>
            <div class="hero-subtitle">
                Responses are returned from matching text inside your uploaded documents only.
                When the information is missing, the assistant returns <strong>{ui_text(language, "not_found")}</strong>.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_empty_state(documents: list[dict[str, Any]]) -> None:
    if documents:
        preview_columns = st.columns(min(3, len(documents)))
        for column, document in zip(preview_columns, documents[:3]):
            with column:
                st.markdown(
                    f"""
                    <div class="panel-card">
                        <strong>{document["original_name"]}</strong><br>
                        <span class="helper-text">{document["page_count"]} pages - {document["chunk_count"]} chunks</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
    else:
        st.info("Upload PDFs from the sidebar to create a private knowledge base.")


def render_voice_tools(provider: str, language: str) -> None:
    st.markdown('<div class="panel-card">', unsafe_allow_html=True)
    st.markdown("**Voice assistant**")
    st.caption("Record a question, transcribe it offline, and send it into the current chat.")
    audio_value = st.audio_input("Ask by voice")
    if st.button("Transcribe voice question", use_container_width=True):
        if not audio_value:
            st.warning("Record audio first.")
        else:
            try:
                ai_service = get_ai_service(provider)
                transcript = ai_service.transcribe_audio(audio_value.getvalue())
                if transcript:
                    st.session_state.pending_prompt = transcript
                    st.success("Voice question transcribed. It is ready in the chat box.")
                    st.rerun()
                else:
                    st.warning("No speech was detected in the recording.")
            except Exception as exc:
                st.error(str(exc))
    st.markdown(f"**{ui_text(language, 'voice_output')}**")
    st.caption(ui_text(language, "voice_output_help"))
    last_answer = st.session_state.get("last_answer_to_speak", "")
    if last_answer:
        render_speech_player(last_answer, LANGUAGE_CONFIG[language]["speech"], language)
    else:
        st.info(ui_text(language, "voice_output_empty"))
    st.markdown("</div>", unsafe_allow_html=True)


def render_scan_import_tools(user_id: int, provider: str) -> int:
    st.markdown("### Scan or Import")
    st.caption("Capture pages with your camera or choose images from gallery/storage, then turn them into a searchable PDF.")
    camera_tab, gallery_tab = st.tabs(["📷 Camera Scan", "🖼️ Gallery Upload"])

    with camera_tab:
        camera_controls = st.columns(2)
        with camera_controls[0]:
            open_camera = st.button("📷 Open Camera", key="open_camera_scan", use_container_width=True)
        with camera_controls[1]:
            stop_camera = st.button(
                "⏹️ Stop Camera",
                key="stop_camera_scan",
                use_container_width=True,
                disabled=not st.session_state.get("camera_scan_enabled", False),
            )

        if open_camera and not st.session_state.get("camera_scan_enabled", False):
            st.session_state.camera_scan_enabled = True
            st.session_state.camera_scan_input_key = int(st.session_state.get("camera_scan_input_key", 0)) + 1
            st.rerun()

        if stop_camera:
            st.session_state.camera_scan_enabled = False
            st.session_state.camera_scan_input_key = int(st.session_state.get("camera_scan_input_key", 0)) + 1
            st.rerun()

        camera_capture = None
        if st.session_state.get("camera_scan_enabled", False):
            st.caption("Camera access is requested only after you choose to open it.")
            camera_capture = st.camera_input(
                "Scan document",
                key=f"camera_scan_input_{st.session_state.get('camera_scan_input_key', 0)}",
            )
            if camera_capture:
                st.image(camera_capture.getvalue(), caption="Current camera capture", use_container_width=True)
        else:
            st.info("Camera is off. Use Open Camera to start scanning pages manually.")

        if st.button("📄 Scan Document", key="add_camera_page", use_container_width=True):
            if not st.session_state.get("camera_scan_enabled", False):
                st.warning("Open the camera first to scan a page.")
            elif not camera_capture:
                st.warning("Capture a page first.")
            else:
                page_bytes = camera_capture.getvalue()
                page_hash = compute_file_hash(page_bytes)
                known_hashes = list(st.session_state.get("camera_scan_hashes", []))
                if page_hash in known_hashes:
                    st.info("That camera page is already in the scan queue.")
                else:
                    queued_pages = list(st.session_state.get("camera_scan_pages", []))
                    queued_pages.append(
                        {
                            "name": f"camera_page_{len(queued_pages) + 1}.jpg",
                            "bytes": page_bytes,
                        }
                    )
                    known_hashes.append(page_hash)
                    st.session_state.camera_scan_pages = queued_pages
                    st.session_state.camera_scan_hashes = known_hashes
                    st.success(f"Added camera page {len(queued_pages)} to the scan queue.")

    with gallery_tab:
        gallery_files = st.file_uploader(
            "Choose images from gallery or storage",
            type=IMAGE_UPLOAD_TYPES,
            accept_multiple_files=True,
            key="gallery_scan_upload",
        )
        if gallery_files:
            st.image([file.getvalue() for file in gallery_files], caption=[file.name for file in gallery_files], use_container_width=True)

    queued_camera_pages = list(st.session_state.get("camera_scan_pages", []))
    if queued_camera_pages:
        st.caption(f"Camera queue: {len(queued_camera_pages)} page(s)")
        st.image([page["bytes"] for page in queued_camera_pages], caption=[page["name"] for page in queued_camera_pages], use_container_width=True)

    scan_file_name = st.text_input("Generated PDF name", value="scanned_notes.pdf", key="scan_pdf_name")
    action_cols = st.columns(2)
    with action_cols[0]:
        process_scan = st.button("📄 Create searchable PDF", key="process_scan_pages", use_container_width=True)
    with action_cols[1]:
        clear_scan = st.button("Clear scan queue", key="clear_scan_queue", use_container_width=True)

    if clear_scan:
        reset_scan_state()
        st.rerun()

    if not ocr_is_available():
        st.info("OCR for camera/gallery scans is ready once `pytesseract` and the Tesseract OCR app are installed.")

    processed = 0
    if process_scan:
        try:
            ai_service = get_ai_service(provider)
            generated_file, extracted_text, page_count, page_records, preview_images = build_scanned_document(
                queued_camera_pages,
                gallery_files or [],
                scan_file_name,
            )
            st.image(preview_images, caption=[f"Enhanced page {index}" for index in range(1, len(preview_images) + 1)], use_container_width=True)
            ok, message = store_document_bytes(
                user_id,
                generated_file,
                ai_service,
                generated_file.getvalue(),
                extracted_text,
                page_count,
                page_records,
            )
            st.success(message) if ok else st.warning(message)
            processed = int(ok)
            if ok:
                reset_scan_state()
                st.rerun()
        except Exception as exc:
            st.error(str(exc))

    return processed


def render_chat_messages(messages: list[dict[str, Any]]) -> None:
    not_found_values = {TEXTS["not_found"][config["code"]] for config in LANGUAGE_CONFIG.values()}
    not_found_values.add(INFO_NOT_FOUND)
    for message in messages:
        role = "user" if message["role"] == "user" else "assistant"
        bubble_class = "user-bubble" if role == "user" else "assistant-bubble"
        with st.chat_message(role):
            st.markdown(f'<div class="{bubble_class}">{message["content"]}</div>', unsafe_allow_html=True)
            metadata = message.get("metadata") or {}
            sources = metadata.get("sources") or []
            if sources and role == "assistant" and message["content"] not in not_found_values:
                st.caption(f"Sources: {', '.join(sources)}")


def render_prompt_box(language: str) -> str | None:
    pending = st.session_state.get("pending_prompt", "")
    if pending:
        st.info(f"{ui_text(language, 'voice_ready')}: {pending}")
        if st.button("Use voice draft as next prompt", use_container_width=False):
            st.session_state.pending_prompt = ""
            return pending

    prompt = st.chat_input("Ask a question from the selected PDFs only")
    return prompt


def render_export_tools(messages: list[dict[str, Any]], title: str, language: str) -> None:
    if not messages:
        st.info("Start a chat to unlock exports.")
        return
    st.markdown(f"**{ui_text(language, 'export')}**")
    export_cols = st.columns(3)
    with export_cols[0]:
        st.download_button(
            "PDF",
            data=build_chat_pdf(messages, title),
            file_name=f"{clean_filename(title)}.pdf",
            mime="application/pdf",
            use_container_width=True,
        )
    with export_cols[1]:
        st.download_button(
            "TXT",
            data=build_chat_txt(messages, title),
            file_name=f"{clean_filename(title)}.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with export_cols[2]:
        if DocxFile is None:
            st.button("DOCX unavailable", disabled=True, use_container_width=True)
        else:
            st.download_button(
                "DOCX",
                data=build_chat_docx(messages, title),
                file_name=f"{clean_filename(title)}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )


def render_study_tools(selected_docs: list[dict[str, Any]], language: str) -> None:
    st.markdown(f"### {ui_text(language, 'study_tools')}")
    if not selected_docs:
        st.info("Select at least one PDF to generate study material.")
        return

    study_tabs = st.tabs(
        [
            ui_text(language, "summary"),
            ui_text(language, "questions"),
            ui_text(language, "mcq"),
            ui_text(language, "notes"),
        ]
    )

    with study_tabs[0]:
        summary = build_summary(selected_docs)
        if not summary:
            st.info(ui_text(language, "not_found"))
        for line in format_study_lines(summary):
            st.markdown(line)

    with study_tabs[1]:
        questions = build_important_questions(selected_docs)
        if not questions:
            st.info(ui_text(language, "not_found"))
        for line in format_study_lines(questions, include_question=True):
            st.markdown(line)

    with study_tabs[2]:
        mcqs = build_mcqs(selected_docs)
        if not mcqs:
            st.info(ui_text(language, "not_found"))
        for index, mcq in enumerate(mcqs, start=1):
            st.markdown(f"**{index}. {mcq['question']}**")
            for option_index, option in enumerate(mcq["options"][:4]):
                label = chr(ord("A") + option_index)
                st.markdown(f"{label}. {option}")
            st.caption(f"Answer: {mcq['answer']} | {mcq['source']} - {page_label(mcq.get('page'))}")

    with study_tabs[3]:
        notes = build_notes(selected_docs)
        if not notes:
            st.info(ui_text(language, "not_found"))
        for line in format_study_lines(notes):
            st.markdown(line)


def render_page_search(selected_docs: list[dict[str, Any]], provider: str, language: str) -> None:
    st.markdown(f"### {ui_text(language, 'page_search_title')}")
    st.caption(ui_text(language, "page_search_help"))
    if not selected_docs:
        st.info("Select at least one PDF to search.")
        return

    query = st.text_input(
        ui_text(language, "search_query"),
        value=st.session_state.get("page_search_query", ""),
        key="page_search_box",
    )
    st.session_state.page_search_query = query
    if st.button(ui_text(language, "search_button"), use_container_width=False):
        if not query.strip():
            st.warning("Enter a search query first.")
            return
        ai_service = get_ai_service(provider)
        _, _, _, matches = retrieve_relevant_context(ai_service, selected_docs, query)
        if not matches:
            st.warning(ui_text(language, "not_found"))
        for match in matches:
            st.markdown(
                f"""
                <div class="panel-card" style="margin-bottom: 0.75rem;">
                    <strong>{match["source"]}</strong><br>
                    <span class="helper-text">{page_label(match.get("page"))} · Score {match["score"]}</span>
                    <div style="margin-top: 0.45rem;">{match["excerpt"]}</div>
                </div>
                """,
                unsafe_allow_html=True,
            )


def render_analytics_dashboard(user: dict[str, Any], analytics: dict[str, int], language: str) -> None:
    st.markdown(f"### {ui_text(language, 'dashboard')}")
    metric_values = [
        ("PDF count", analytics["pdf_count"]),
        ("Chats", analytics["chat_count"]),
        ("Questions", analytics["question_count"]),
        ("Pages", analytics["page_count"]),
        ("Chunks", analytics["chunk_count"]),
        ("Cloud PDFs", analytics["cloud_count"]),
    ]
    for start in range(0, len(metric_values), 3):
        metric_cols = st.columns(3)
        for column, (label, value) in zip(metric_cols, metric_values[start : start + 3]):
            with column:
                st.markdown(
                    f"""
                    <div class="metric-card">
                        <div class="helper-text">{label}</div>
                        <div class="metric-value">{value}</div>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )

    st.markdown(
        f"""
        <div class="panel-card" style="margin-top: 1rem;">
            <strong>{ui_text(language, 'analytics')}</strong><br>
            <span class="helper-text">
                Account: {user["email"]}<br>
                Cloud mirror directory: {CLOUD_DIR}
            </span>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_chat_workspace(user: dict[str, Any], documents: list[dict[str, Any]], selected_doc_ids: list[int], provider: str, language: str) -> None:
    chats = get_user_chats(user["id"])
    analytics = get_user_analytics(user["id"])
    render_workspace_header(documents, chats, provider, language, analytics)

    document_lookup = {doc["id"]: doc for doc in documents}
    active_chat = get_chat(st.session_state.active_chat_id, user["id"]) if st.session_state.active_chat_id else None
    selected_docs = [document_lookup[doc_id] for doc_id in selected_doc_ids if doc_id in document_lookup]
    if selected_docs:
        try:
            selected_docs = ensure_documents_ready(selected_docs, get_ai_service(provider))
        except Exception as exc:
            st.error(str(exc))
            return
    workspace_tabs = st.tabs(
        [
            ui_text(language, "chat"),
            ui_text(language, "study_tools"),
            ui_text(language, "smart_search"),
            ui_text(language, "dashboard"),
        ]
    )

    with workspace_tabs[0]:
        left_col, right_col = st.columns([1.9, 0.9], gap="large")

        with left_col:
            st.markdown('<div class="chat-shell">', unsafe_allow_html=True)
            if active_chat:
                st.markdown(
                    f"""
                    <div class="panel-card" style="margin-bottom: 1rem;">
                        <strong>{active_chat["title"]}</strong><br>
                        <span class="helper-text">Mode: Offline PDF search</span>
                    </div>
                    """,
                    unsafe_allow_html=True,
                )
                messages = get_messages(active_chat["id"])
                render_chat_messages(messages)
            else:
                messages = []
                render_empty_state(documents)

            prompt = render_prompt_box(language)
            st.markdown("</div>", unsafe_allow_html=True)

        with right_col:
            render_voice_tools(provider, language)
            render_export_tools(messages, active_chat["title"] if active_chat else "chat", language)
            if active_chat:
                if st.button(f"🗑️ {ui_text(language, 'clear_current_chat')}", key="chat_panel_clear", use_container_width=True):
                    delete_chat(active_chat["id"], user["id"])
                    st.session_state.active_chat_id = None
                    st.session_state.pending_prompt = ""
                    st.session_state.last_answer_to_speak = ""
                    reset_chat_clear_state()
                    st.rerun()
            st.markdown(
                """
                <div class="panel-card">
                    <strong>Deployment notes</strong><br>
                    <span class="helper-text">
                        Set <code>DATABASE_URL</code> for PostgreSQL in production. PDF uploads are mirrored
                        into the configured cloud storage directory and remain available for offline study tools.
                    </span>
                </div>
                """,
                unsafe_allow_html=True,
            )

    with workspace_tabs[1]:
        render_study_tools(selected_docs, language)

    with workspace_tabs[2]:
        render_page_search(selected_docs, provider, language)

    with workspace_tabs[3]:
        render_analytics_dashboard(user, analytics, language)

    if not prompt:
        return

    if not selected_doc_ids:
        st.warning("Select at least one PDF before asking a question.")
        return

    if not selected_docs:
        st.warning("The selected documents are no longer available.")
        return

    ai_service = get_ai_service(provider)

    if not active_chat:
        chat_title = build_chat_title(prompt, selected_docs)
        chat_id = create_chat(user["id"], chat_title, selected_doc_ids, provider)
        st.session_state.active_chat_id = chat_id
        active_chat = get_chat(chat_id, user["id"])

    add_message(
        active_chat["id"],
        "user",
        prompt,
        metadata={"selected_document_ids": selected_doc_ids, "provider": provider},
    )
    with st.chat_message("user"):
        st.markdown(f'<div class="user-bubble">{prompt}</div>', unsafe_allow_html=True)

    with st.chat_message("assistant"):
        answer_placeholder = st.empty()
        source_placeholder = st.empty()
        answer_placeholder.markdown(f'<div class="assistant-bubble">Thinking...</div>', unsafe_allow_html=True)
        try:
            contexts, sources, best_score, matches = retrieve_relevant_context(ai_service, selected_docs, prompt)
            answer = ai_service.generate_answer(prompt, contexts)
            if answer == INFO_NOT_FOUND:
                answer = ui_text(language, "not_found")
            answer_placeholder.markdown(f'<div class="assistant-bubble">{answer}</div>', unsafe_allow_html=True)
            if sources and answer != ui_text(language, "not_found"):
                source_placeholder.caption(f"Sources: {', '.join(sources)}")
        except Exception as exc:
            answer = f"Error: {exc}"
            answer_placeholder.markdown(f'<div class="assistant-bubble">{answer}</div>', unsafe_allow_html=True)
            sources = []
            best_score = 0.0
            matches = []

    st.session_state.last_answer_to_speak = answer if answer != ui_text(language, "not_found") else ""
    add_message(
        active_chat["id"],
        "assistant",
        answer,
        metadata={
            "sources": sources,
            "retrieval_score": best_score,
            "provider": provider,
            "page_results": matches,
            "matched_contexts": len(contexts) if "contexts" in locals() else 0,
        },
    )
    st.rerun()


def main() -> None:
    init_db()
    init_session_state()

    user = st.session_state.user
    if user:
        st.session_state.theme = user.get("theme_preference", st.session_state.theme)
    inject_styles(st.session_state.theme)

    if not user:
        render_auth_screen()
        return

    documents, selected_doc_ids, provider, language = render_sidebar(user)
    render_chat_workspace(user, documents, selected_doc_ids, provider, language)


if __name__ == "__main__":
    main()
